import os
from typing import Dict, Optional, List
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image


def detect_document_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    type_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".xlsx": "xlsx",
        ".xls": "xls",
        ".txt": "txt",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".gif": "image",
        ".bmp": "image",
    }
    return type_map.get(ext, "unknown")


def extract_pdf_metadata(file_path: str) -> Dict:
    metadata = {}
    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            info = reader.metadata
            if info:
                metadata["title"] = info.title or ""
                metadata["author"] = info.author or ""
                metadata["subject"] = info.subject or ""
                metadata["creator"] = info.creator or ""
                metadata["producer"] = info.producer or ""
                metadata["num_pages"] = len(reader.pages)
    except Exception as e:
        metadata["error"] = str(e)
    return metadata


def extract_docx_metadata(file_path: str) -> Dict:
    metadata = {}
    try:
        doc = DocxDocument(file_path)
        core_props = doc.core_properties
        metadata["title"] = core_props.title or ""
        metadata["author"] = core_props.author or ""
        metadata["subject"] = core_props.subject or ""
        metadata["created"] = str(core_props.created) if core_props.created else ""
        metadata["modified"] = str(core_props.modified) if core_props.modified else ""
    except Exception as e:
        metadata["error"] = str(e)
    return metadata


def extract_xlsx_metadata(file_path: str) -> Dict:
    metadata = {}
    try:
        wb = load_workbook(file_path, read_only=True)
        metadata["sheet_names"] = wb.sheetnames
        metadata["num_sheets"] = len(wb.sheetnames)
        wb.close()
    except Exception as e:
        metadata["error"] = str(e)
    return metadata


def extract_image_metadata(file_path: str) -> Dict:
    metadata = {}
    try:
        with Image.open(file_path) as img:
            metadata["format"] = img.format
            metadata["size"] = img.size
            metadata["mode"] = img.mode
    except Exception as e:
        metadata["error"] = str(e)
    return metadata


def extract_document_metadata(file_path: str, doc_type: str) -> Dict:
    metadata = {
        "file_size": os.path.getsize(file_path),
        "file_name": os.path.basename(file_path),
        "file_type": doc_type,
    }

    if doc_type == "pdf":
        metadata.update(extract_pdf_metadata(file_path))
    elif doc_type == "docx":
        metadata.update(extract_docx_metadata(file_path))
    elif doc_type == "xlsx":
        metadata.update(extract_xlsx_metadata(file_path))
    elif doc_type == "image":
        metadata.update(extract_image_metadata(file_path))

    return metadata


def extract_pdf_text(file_path: str) -> str:
    text = ""
    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
    except Exception:
        pass
    return text


def extract_docx_text(file_path: str) -> str:
    text = ""
    try:
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception:
        pass
    return text


def extract_txt_text(file_path: str) -> str:
    text = ""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except Exception:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
        except Exception:
            pass
    return text


def extract_document_text(file_path: str, doc_type: str) -> str:
    if doc_type == "pdf":
        return extract_pdf_text(file_path)
    elif doc_type == "docx":
        return extract_docx_text(file_path)
    elif doc_type == "txt":
        return extract_txt_text(file_path)
    else:
        return ""
